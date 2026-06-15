# -*- coding: utf-8 -*-
"""
construir_reporte_docx.py
=========================
Genera el documento Word de analisis comparativo de modelos con python-docx.
Salida: docs/Analisis_Comparativo_Modelos.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = "1F4E79"
AZUL_CLARO = "D6E4F0"
GRIS = "F2F2F2"


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=None, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, w in enumerate(widths_inches):
            row.cells[i].width = Inches(w)


def add_table(doc, encabezados, filas, widths, header_size=9, body_size=8.5):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(encabezados):
        c = t.rows[0].cells[i]
        set_cell_text(c, h, bold=True, size=header_size, color="FFFFFF", align="center")
        set_cell_background(c, AZUL)
    for r, fila in enumerate(filas):
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            set_cell_text(cells[i], val, size=body_size,
                          align="center" if i > 0 else "left")
        if r % 2 == 1:
            for c in cells:
                set_cell_background(c, GRIS)
    set_col_widths(t, widths)
    return t


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


# ===========================================================================
doc = Document()

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# Margenes / tamano carta
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin = sec.bottom_margin = Inches(1)

# ---- Portada ----
titulo = doc.add_heading("Análisis comparativo de modelos preentrenados de Hugging Face", level=0)
sub = doc.add_paragraph()
r = sub.add_run("Caso de uso: análisis de sentimiento en español (NLP)")
r.bold = True
r.font.size = Pt(14)
doc.add_paragraph("Reporte técnico — Selección e implementación de modelos preentrenados")
meta = doc.add_paragraph()
meta.add_run("Entorno de prototipado: ").bold = True
meta.add_run("Google Colab (GPU NVIDIA T4)   |   ")
meta.add_run("Ecosistema: ").bold = True
meta.add_run("Hugging Face Hub (Models, Datasets, Spaces)")
doc.add_paragraph("Fecha: junio de 2026")

doc.add_paragraph()

# ---- Resumen ejecutivo ----
h(doc, "1. Resumen ejecutivo", 1)
doc.add_paragraph(
    "Este reporte documenta la selección, implementación y comparación de tres modelos "
    "preentrenados de clasificación de sentimiento en español obtenidos de Hugging Face Hub. "
    "Los modelos se evaluaron sobre un conjunto balanceado de 120 reseñas/opiniones cortas "
    "(clases POS, NEU, NEG) midiendo accuracy, precision, recall y F1-score (macro), así como "
    "la latencia de inferencia. El objetivo es identificar el modelo más adecuado equilibrando "
    "precisión, eficiencia, licencia y recursos."
)
doc.add_paragraph(
    "Conclusión principal: para un prototipo académico centrado en español, se recomienda "
    "RoBERTuito (pysentimiento/robertuito-sentiment-analysis) por su mejor relación "
    "precisión/recursos en el idioma. Si el despliegue requiere uso comercial o cobertura "
    "multilingüe, las alternativas recomendadas son XLM-RoBERTa de Cardiff NLP (mayor robustez) "
    "o el DistilBERT multilingüe destilado (máxima eficiencia y licencia Apache-2.0)."
)

# ---- Caso de uso ----
h(doc, "2. Descripción del caso de uso", 1)
doc.add_paragraph(
    "Se desea incorporar un componente de análisis de sentimiento a un sistema que procesa "
    "opiniones de clientes en español (reseñas de productos, comentarios de servicio y "
    "menciones en redes sociales). El sistema debe clasificar cada texto en una de tres "
    "categorías —positivo (POS), neutral (NEU) o negativo (NEG)— para alimentar tableros de "
    "voz del cliente y priorizar la atención de comentarios negativos."
)
doc.add_paragraph("Requisitos del componente:")
add_bullets(doc, [
    "Idioma principal: español (incluye texto corto e informal de redes sociales).",
    "Tres clases de salida: POS / NEU / NEG.",
    "Inferencia rápida para escenarios interactivos y de alto volumen.",
    "Ejecutable en recursos modestos (una GPU T4 de Colab, idealmente también en CPU).",
    "Preferible no requerir entrenamiento propio: uso de modelos preentrenados (zero-shot).",
])

h(doc, "2.1 Componentes del ecosistema Hugging Face utilizados", 2)
add_bullets(doc, [
    "Models: los tres modelos preentrenados de clasificación de texto descargados del Hub.",
    "Datasets: referencia a corpus estándar (TASS 2020, InterTASS 2017, multilingual-sentiments) "
    "usados para los benchmarks publicados; el dataset de evaluación propio se versiona en el repositorio.",
    "Spaces: demos interactivas de estos modelos en el Hub, útiles para validación cualitativa previa.",
])

# ---- Criterios de seleccion ----
h(doc, "3. Criterios de selección", 1)
doc.add_paragraph(
    "La selección combinó criterios técnicos (calidad y eficiencia del modelo) y prácticos "
    "(licencia, madurez y compatibilidad). Se ponderó así:"
)
add_table(
    doc,
    ["Criterio", "Descripción", "Prioridad"],
    [
        ["Idioma y dominio", "Desempeño en español, en especial texto corto/informal.", "Alta"],
        ["Precisión (F1 macro)", "Calidad de clasificación en las tres clases.", "Alta"],
        ["Latencia / eficiencia", "Tiempo de inferencia por muestra (interactivo y por lote).", "Media-Alta"],
        ["Tamaño y recursos", "Parámetros, memoria y peso de descarga; cabe en T4/CPU.", "Media"],
        ["Licencia y uso permitido", "Compatibilidad legal con el tipo de despliegue.", "Alta"],
        ["Madurez y mantenimiento", "Descargas, comunidad y estabilidad/disponibilidad en el Hub.", "Media"],
        ["Esquema de etiquetas", "Salida directa de 3 clases (POS/NEU/NEG) o mapeable.", "Alta"],
    ],
    widths=[1.6, 3.7, 1.1],
    body_size=9,
)
doc.add_paragraph(
    "Con estos criterios se descartó finiteautomata/beto-sentiment-analysis (BETO): aunque es un "
    "modelo válido en español, su ficha en el Hub advierte que “será eliminado pronto”, lo que "
    "representa un riesgo de reproducibilidad y mantenimiento. Se sustituyó por un DistilBERT "
    "multilingüe destilado, que además aporta un punto de comparación de eficiencia y una licencia permisiva."
)

# ---- Modelos evaluados ----
h(doc, "4. Modelos evaluados", 1)
doc.add_paragraph("Se seleccionaron tres modelos de arquitecturas y perfiles complementarios:")
add_bullets(doc, [
    "RoBERTuito (pysentimiento/robertuito-sentiment-analysis): RoBERTa preentrenado con ~500 M de "
    "tweets en español y ajustado con TASS 2020. Especializado en español informal. 12 capas, ~108 M parámetros.",
    "DistilBERT multilingüe destilado (lxyuan/distilbert-base-multilingual-cased-sentiments-student): "
    "modelo destilado (zero-shot) a partir de mDeBERTa, sobre el dataset multilingual-sentiments (12 idiomas). "
    "Solo 6 capas → el más rápido. ~135 M parámetros. Licencia Apache-2.0.",
    "XLM-RoBERTa Twitter (cardiffnlp/twitter-xlm-roberta-base-sentiment): XLM-R base reentrenado con "
    "~198 M de tweets y ajustado en el benchmark UMSAB (8 idiomas, incl. español). 12 capas, ~278 M parámetros. "
    "El más robusto en escenarios multilingües.",
])

# ---- Metodologia ----
h(doc, "5. Metodología", 1)
add_bullets(doc, [
    "Dataset de evaluación: 120 textos en español, balanceado (40 POS / 40 NEU / 40 NEG), "
    "curado por los autores (ver data/README_dataset.md).",
    "Inferencia: pipeline de text-classification de Transformers, por lotes (batch=16), "
    "truncamiento a 256 tokens, en GPU T4.",
    "Métricas: accuracy y precision/recall/F1 macro (sklearn). El esquema balanceado hace que "
    "accuracy y F1 macro sean directamente interpretables.",
    "Latencia: tiempo medio por muestra con batch=1 (escenario de petición individual), "
    "descartando una pasada de calentamiento; se reporta el mejor de 3 ensayos.",
    "Evaluación zero-shot: se usan los modelos tal cual, sin fine-tuning.",
])
doc.add_paragraph(
    "Nota sobre las cifras de precisión de la Tabla 1: corresponden a benchmarks publicados por los "
    "autores de cada modelo (Macro-F1). El cuaderno de Colab recalcula accuracy, precision, recall, "
    "F1 y latencia reales sobre el dataset del proyecto; esos valores empíricos complementan la tabla "
    "y pueden variar según el hardware y la versión de las librerías."
)

# ---- Tabla comparativa ----
h(doc, "6. Tabla comparativa de modelos", 1)
doc.add_paragraph("Tabla 1. Comparativa de los tres modelos según los criterios definidos.")
add_table(
    doc,
    ["Modelo", "Tarea", "Tamaño", "Licencia", "Precisión (Macro-F1, referencia)",
     "Latencia típica (T4)", "Costo / Recursos"],
    [
        ["RoBERTuito\n(robertuito-sentiment-analysis)", "Sent. ES (3 clases)", "~108 M\n(12 capas)",
         "No comercial / investigación", "0.705 — TASS 2020", "Media (~5–9 ms/muestra)",
         "Baja: ~0.43 GB; corre en T4 y CPU"],
        ["DistilBERT-multi\n(sentiments-student)", "Sent. multiling. (3 clases)", "~135 M\n(6 capas)",
         "Apache-2.0", "No publicado (destilado); estimado menor", "Muy baja (~2–5 ms/muestra)",
         "Muy baja: ~0.54 GB; ideal alto volumen/edge"],
        ["XLM-RoBERTa Twitter\n(twitter-xlm-roberta-base-sentiment)", "Sent. multiling. (3 clases)",
         "~278 M\n(12 capas)", "Académica (XLM-T, MIT)", "0.679 — InterTASS 2017 (UMSAB)",
         "Alta (~8–15 ms/muestra)", "Media: ~1.1 GB; mayor memoria/cómputo"],
    ],
    widths=[1.35, 0.95, 0.7, 0.95, 1.15, 0.95, 1.3],
    header_size=8.5,
    body_size=8,
)
doc.add_paragraph(
    "Notas: (1) Las cifras de Macro-F1 provienen de conjuntos de prueba distintos pero comparables "
    "(TASS 2020 e InterTASS 2017, ambos sentimiento de tweets en español, 3 clases). El DistilBERT "
    "destilado no reporta F1 de sentimiento estándar; su ficha indica 88.3 % de acuerdo con el modelo "
    "maestro. (2) Las latencias son aproximadas para una GPU T4 con inferencia por lote; en CPU son "
    "~10–30× mayores. El cuaderno mide la latencia real del entorno."
)

# ---- Trade-offs ----
h(doc, "7. Análisis de trade-offs (eficiencia frente a precisión)", 1)
doc.add_paragraph(
    "El conjunto de modelos forma un claro espacio de compromiso entre precisión y eficiencia:"
)
add_bullets(doc, [
    "Precisión en español: RoBERTuito lidera (0.705) por estar especializado en el idioma y dominio; "
    "XLM-RoBERTa queda cerca (0.679) aportando robustez multilingüe; el DistilBERT destilado, al "
    "derivar de destilación zero-shot, suele rendir por debajo de ambos.",
    "Eficiencia: el factor decisivo no es solo el número de parámetros, sino la profundidad. "
    "DistilBERT tiene ~135 M parámetros (más que RoBERTuito), pero solo 6 capas de Transformer "
    "frente a 12; como el cómputo lo dominan las capas (no la tabla de embeddings), DistilBERT es "
    "el más rápido. XLM-R, con 12 capas y un vocabulario de 250 k, es el más costoso.",
    "Costo/recursos: los tres caben holgadamente en una T4 (16 GB). XLM-R exige ~2.5× el peso de "
    "RoBERTuito en disco y memoria, lo que importa para despliegues en CPU o a gran escala.",
    "Licencia: solo DistilBERT (Apache-2.0) permite uso comercial sin fricción; RoBERTuito es de uso "
    "no comercial/investigación; XLM-R hereda restricciones de los datasets de origen.",
])
doc.add_paragraph(
    "En síntesis: RoBERTuito ofrece la mejor precisión por unidad de recurso en español; DistilBERT "
    "maximiza throughput y flexibilidad de licencia a costa de precisión; XLM-R es la apuesta por "
    "robustez multilingüe cuando el presupuesto de latencia lo permite."
)

# ---- Decisiones de ajuste minimo ----
h(doc, "8. Decisiones de ajuste mínimo", 1)
add_bullets(doc, [
    "Normalización de etiquetas: unificar las salidas heterogéneas (POS/NEU/NEG, "
    "positive/neutral/negative, Positive/Neutral/Negative) a un esquema canónico {NEG, NEU, POS}. "
    "Es el único ajuste imprescindible para comparar de forma homogénea.",
    "Preprocesamiento ligero: sustituir menciones (@usuario) y URLs por marcadores, según la "
    "convención de los modelos entrenados con tweets.",
    "Truncamiento a 256 tokens para acotar latencia y memoria sin afectar textos cortos.",
    "Inferencia por lotes (batch=16) para throughput; medición de latencia con batch=1 para el "
    "escenario interactivo.",
    "Sin fine-tuning en esta fase. Mejoras futuras de bajo costo: ajuste de umbral de decisión por "
    "clase (sobre todo NEU), calibración de probabilidades y, si se requiere, fine-tuning ligero con "
    "datos del dominio.",
])

# ---- Justificacion ----
h(doc, "9. Justificación del modelo recomendado", 1)
doc.add_paragraph(
    "Para el prototipo académico orientado a español se recomienda RoBERTuito por:"
)
add_bullets(doc, [
    "Mejor precisión publicada en español (Macro-F1 0.705) entre los candidatos.",
    "Tamaño compacto (~108 M) y latencia media-baja, adecuados para T4 e incluso CPU.",
    "Salida nativa de 3 clases (POS/NEU/NEG), alineada con el caso de uso.",
    "Amplia adopción y madurez en el Hub (más de 600 k descargas mensuales).",
])
doc.add_paragraph("Regla de decisión según el contexto de despliegue:")
add_bullets(doc, [
    "Prototipo / investigación en español → RoBERTuito.",
    "Producción comercial o requisitos de licencia estrictos → DistilBERT (Apache-2.0); "
    "elegirlo también cuando la latencia/escala sea crítica.",
    "Cobertura multilingüe o textos más variados → XLM-RoBERTa, si el presupuesto de latencia lo permite.",
])
doc.add_paragraph(
    "La recomendación final debe confirmarse con los resultados empíricos del cuaderno sobre el "
    "dataset del proyecto, ya que el dominio (reseñas) difiere parcialmente del dominio de "
    "entrenamiento (tweets)."
)

# ---- Umbrales y riesgos ----
h(doc, "10. Umbrales y riesgos", 1)
h(doc, "10.1 Umbrales propuestos", 2)
add_bullets(doc, [
    "Calidad mínima aceptable: Macro-F1 ≥ 0.65 en el dataset de evaluación; por debajo, no desplegar.",
    "Latencia objetivo: ≤ 50 ms/muestra para uso interactivo; ≤ 10 ms/muestra para alto volumen.",
    "Confianza: derivar a revisión humana las predicciones con probabilidad máxima < 0.50, "
    "especialmente las clasificadas como NEU.",
])
h(doc, "10.2 Riesgos identificados", 2)
add_bullets(doc, [
    "Ironía y sarcasmo: los modelos de sentimiento suelen fallar; riesgo alto en redes sociales.",
    "Desajuste de dominio: entrenados en tweets, pueden perder precisión en reseñas largas o formales.",
    "Clase NEU ambigua: es la más subjetiva y concentra la mayor confusión (NEU↔POS/NEG).",
    "Licencia: RoBERTuito es de uso no comercial; desplegarlo comercialmente implica riesgo legal.",
    "Mantenimiento/disponibilidad: los modelos de terceros pueden ser despublicados (caso BETO). "
    "Mitigación: fijar una revisión concreta (commit/hash) del modelo en el Hub y conservar una copia.",
    "Variación dialectal y cambio de código (code-switching) español-inglés.",
    "Tamaño de muestra reducido (120): intervalos de confianza amplios; no extrapolar sin un test mayor "
    "(p. ej., TASS 2020 o amazon_reviews_multi-es).",
])

# ---- Reproducibilidad ----
h(doc, "11. Reproducibilidad y entorno", 1)
add_bullets(doc, [
    "Entorno: Google Colab con GPU T4; Python 3.10+; transformers ≥ 4.40, huggingface_hub ≥ 0.23, "
    "sentencepiece, scikit-learn, pandas, matplotlib.",
    "Repositorio (GitHub): README.md, requirements.txt, data/ (dataset + diccionario), "
    "notebooks/ (cuaderno .ipynb), src/ (scripts) y results/ (tablas y figuras).",
    "Pasos: abrir el cuaderno en Colab → activar GPU T4 → Ejecutar todo. Tiempo aprox. 3–6 min "
    "(incluye descarga de modelos).",
    "El dataset se regenera de forma determinística con src/generar_dataset.py; la evaluación se "
    "reproduce con src/evaluar_modelos.py.",
])

# ---- Conclusiones ----
h(doc, "12. Conclusiones", 1)
doc.add_paragraph(
    "La comparación muestra que no existe un único “mejor” modelo, sino el más adecuado según el "
    "contexto. RoBERTuito ofrece el mejor equilibrio precisión/recursos en español y es la elección "
    "para el prototipo. DistilBERT destaca en eficiencia y licencia, y XLM-RoBERTa en robustez "
    "multilingüe. La metodología —normalización de etiquetas, métricas macro y medición de latencia— "
    "permite tomar la decisión de forma objetiva y reproducible, y los umbrales y riesgos definidos "
    "guían un despliegue responsable."
)

# ---- Referencias ----
h(doc, "13. Referencias", 1)
refs = [
    "Pérez, J. M. et al. (2021). pysentimiento: A Python Toolkit for Sentiment Analysis and SocialNLP "
    "tasks. arXiv:2106.09462.",
    "Pérez, J. M. et al. (2022). RoBERTuito: a pre-trained language model for social media text in "
    "Spanish. LREC 2022.",
    "Barbieri, F., Espinosa Anke, L., Camacho-Collados, J. (2022). XLM-T: Multilingual Language Models "
    "in Twitter for Sentiment Analysis and Beyond. LREC 2022 / arXiv:2104.12250.",
    "Cañete, J. et al. (2020). Spanish Pre-trained BERT Model and Evaluation Data (BETO). PML4DC, ICLR.",
    "García-Vega, M. et al. (2020). Overview of TASS 2020. IberLEF/SEPLN.",
    "Fichas de modelo en Hugging Face Hub: pysentimiento/robertuito-sentiment-analysis; "
    "lxyuan/distilbert-base-multilingual-cased-sentiments-student; "
    "cardiffnlp/twitter-xlm-roberta-base-sentiment.",
]
for ref in refs:
    doc.add_paragraph(ref, style="List Number")

os.makedirs("docs", exist_ok=True)
ruta = "docs/Analisis_Comparativo_Modelos.docx"
doc.save(ruta)


def _parchear_zoom(ruta_docx):
    # python-docx escribe <w:zoom w:val="bestFit"/> sin el atributo "percent"
    # que exige el esquema estricto; lo agregamos para que el archivo valide.
    import re
    import zipfile
    z = zipfile.ZipFile(ruta_docx)
    s = z.read("word/settings.xml").decode("utf-8")
    z.close()
    if "<w:zoom" in s and "w:percent" not in s:
        s = re.sub(r'<w:zoom(?![^>]*w:percent)([^>]*)/>',
                   r'<w:zoom\1 w:percent="100"/>', s)
        tmp = ruta_docx + ".tmp"
        zin = zipfile.ZipFile(ruta_docx, "r")
        zout = zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED)
        for item in zin.infolist():
            data = s.encode("utf-8") if item.filename == "word/settings.xml" else zin.read(item.filename)
            zout.writestr(item, data)
        zin.close()
        zout.close()
        os.replace(tmp, ruta_docx)


_parchear_zoom(ruta)
print("Documento generado:", ruta)
print("Parrafos:", len(doc.paragraphs), "| Tablas:", len(doc.tables))
